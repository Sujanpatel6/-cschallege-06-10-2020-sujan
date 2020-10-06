from spellchecker import SpellChecker
import re
import docx
from docx.enum.text import WD_LINE_SPACING

spell = SpellChecker()

# Read a txt file
line=[]
with open("assign_data.txt","r", encoding='utf8') as f:
	for i in f:
		print(i)
		
		line.append(i.strip())

line=[i for i in line if i]
print(line)
# find those words that may be misspelled
doc = docx.Document()


for w in line:
	p = doc.add_paragraph()
	p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
	w=re.sub(" [\(\[].*?[\)\]]","", w)
	w=re.sub("[A-Z\.]{2,}s?","",w)
	
	misspelled = w.split(" ")
	for word in misspelled:
		an_word=''
		for c in word:
			if c.isalnum():
				an_word +=c
			if(c=="."):
				break;
		cword=spell.correction(an_word)
		print(cword)
		if(cword!=an_word):
			runner = p.add_run(cword)
			runner.bold = True
			
		else:
			runner = p.add_run(an_word)
			runner.bold = False
		runner = p.add_run(" ")
			

doc.save('test9.docx')
		
		
